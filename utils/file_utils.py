import io
import docx
import fitz
import json
import base64
import chardet
import pdfplumber
import pandas as pd
from PIL import Image
from flask import current_app
import tiktoken


class FileOperation:

    @staticmethod
    def extract_images_from_pdf(pdf_path, is_content):
        if not is_content:
            return []
        with fitz.open(stream=pdf_path) as doc:
            base64_img = []
            for page_index in range(len(doc)):
                if page_index not in is_content:
                    continue

                page = doc[page_index]
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                img_b64 = base64.b64encode(img_bytes).decode()
                base64_img.append(img_b64)

        return base64_img

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        pdf_bytes = io.BytesIO(pdf_path)
        is_content = []
        final_text = []
        tables = []
        df = ""
        with pdfplumber.open(pdf_bytes) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table and any(any(once) for once in table):
                    tables.append(table)
                    table_bboxes = [pos.bbox for pos in page.find_tables()]
                    filtered_text = page.extract_words()
                    for word in filtered_text:
                        x0, y0, x1, y1 = word["x0"], word["top"], word["x1"], word["bottom"]
                        inside_table = any(
                            t_x0 <= x0 <= t_x1 and t_y0 <= y0 <= t_y1
                            for (t_x0, t_y0, t_x1, t_y1) in table_bboxes
                        )
                        if not inside_table:
                            final_text.append(word["text"])
                else:
                    text = page.extract_text()
                    if text and len(text) > 10:
                        final_text.append(text)
                    else:
                        is_content.append(page.page_number - 1)

            if tables:
                try:
                    df = "\n".join(pd.DataFrame(table[1:], columns=table[0]).to_json(force_ascii=False) for table in tables)
                except:
                    df = json.dumps(tables)

            if final_text:
                final_text = " ".join(final_text) + "\n"
            else:
                final_text = ""

        return final_text + df, is_content

    @staticmethod
    def extract_text_from_word(docx_path):
        doc = docx.Document(docx_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        tables = []
        df = ""

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        if tables:
            try:
                df = pd.DataFrame(tables[0][1:], columns=tables[0][0]).to_json(force_ascii=False)
            except:
                df = json.dumps(tables, ensure_ascii=False)

        return text.strip() + df

    @staticmethod
    def extract_images_from_word(docx_path):
        doc = docx.Document(docx_path)
        base64_img = []
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image = doc.part.rels[rel].target_part.blob
                img = Image.open(io.BytesIO(image))
                buffered = io.BytesIO()
                img.save(buffered, format="PNG")
                img_b64 = base64.b64encode(buffered.getvalue()).decode()
                base64_img.append(img_b64)
        return base64_img

    @staticmethod
    def check_pdf(pdf_path):
        with fitz.open(stream=pdf_path) as doc:
            for page_index in range(len(doc)):
                if doc[page_index].get_images(full=True) or not doc[page_index].get_text("text"):
                    page = doc[page_index]
                    pix = page.get_pixmap(dpi=500)
                    rect = page.rect
                    page.clean_contents()
                    page.insert_image(rect, pixmap=pix)
            pdf_buffer = io.BytesIO()
            doc.save(pdf_buffer, garbage=4, deflate=True)
            pdf_byte_path = pdf_buffer.getvalue()
        return pdf_byte_path

    # -------------------------
    #  修复图片文件名丢失的关键函数
    # -------------------------
    @staticmethod
    def extract_picture(file_stream, file_name):
        """读取图片流并返回 base64 + 文件名"""
        try:
            file_stream.seek(0)
            raw = file_stream.read()
            img_b64 = base64.b64encode(raw).decode()

            return {
                "images": [img_b64],       # 保持 images 仅 base64，不破坏 AI 输入结构
                "filenames": [file_name]   # 修复：统一使用 filenames
            }
        except Exception as e:
            return {
                "images": [],
                "filenames": [file_name],
                "error": str(e)
            }

    # -------------------------
    #        主入口
    # -------------------------
    def __call__(self, username: str, attachment_names: list):
        if not isinstance(attachment_names, list):
            return {"message": "Invalid attachment_names format", "status": 400}
        
        results = {}
        for attachment_name in attachment_names:
            try:
                file_extension = attachment_name.rsplit(".", 1)[1].lower()
                blob_client = current_app.container_client.get_blob_client(f"{username}/{attachment_name}")
                stream = blob_client.download_blob().readall()
                file_stream = io.BytesIO(stream)
                encoding = chardet.detect(stream)["encoding"]
                # -------- TXT --------
                if file_extension == "txt":
                    results[attachment_name] = {
                        "text": stream.decode(encoding),
                        "images": [],
                        "filenames": [attachment_name]
                    }

                # -------- CSV --------
                elif file_extension == "csv":
                    df = pd.read_csv(file_stream, encoding=encoding)
                    results[attachment_name] = {
                        "text": df.to_json(force_ascii=False),
                        "images": [],
                        "filenames": [attachment_name]
                    }

                # -------- JSON --------
                elif file_extension == "json":
                    results[attachment_name] = {
                        "text": stream.decode(encoding),
                        "images": [],
                        "filenames": [attachment_name]
                    }

                # -------- Excel --------
                elif file_extension in ["xlsx", "xls"]:
                    try:
                        if file_extension == "xlsx":
                            df = pd.read_excel(file_stream, engine="openpyxl")
                        else:
                            df = pd.read_excel(file_stream, engine="xlrd")
                        results[attachment_name] = {
                            "text": df.to_json(force_ascii=False),
                            "images": [],
                            "filenames": [attachment_name]     # ⬅️ 修复 EXCEL 也加 filenames
                        }
                    except Exception as e:
                        results[attachment_name] = {
                            "text": f"Failed to read Excel file: {str(e)}",
                            "images": [],
                            "filenames": [attachment_name]
                        }

                # -------- PDF --------
                elif file_extension == "pdf":
                    pdf_stream = io.BytesIO(stream)
                    pdf_text, page_num = self.extract_text_from_pdf(stream)
                    pdf_images = self.extract_images_from_pdf(pdf_stream, page_num)

                    # ⬅️ PDF 多图 → 仍然只有 1 个文件名
                    results[attachment_name] = {
                        "text": pdf_text,
                        "images": pdf_images,
                        "filenames": [attachment_name]
                    }

                # -------- Word DOCX --------
                elif file_extension == "docx":
                    word_text = self.extract_text_from_word(file_stream)
                    try:
                        word_images = self.extract_images_from_word(file_stream)
                    except:
                        word_images = []

                    # ⬅️ Word 多图 → 仍然只有 1 个文件名（文件是一个）
                    results[attachment_name] = {
                        "text": word_text,
                        "images": word_images,
                        "filenames": [attachment_name]
                    }

                # -------- 图片 JPG/PNG --------
                elif file_extension in ["jpg", "jpeg", "png"]:
                    pic = FileOperation.extract_picture(file_stream, attachment_name)
                    print(f"🔍 DEBUG: PNG处理结果 - 文件名: {attachment_name}, filenames: {pic.get('filenames', [])}")

                    # ⬅️ 完整统一格式：text + images + filenames
                    # ⬅️ 支持多张图片拼接 filenames（即便未来支持多图片上传）
                    filenames = pic["filenames"]
                    if isinstance(filenames, str):
                        filenames = [filenames]

                    results[attachment_name] = {
                        "text": f"图片文件: {attachment_name}",
                        "images": pic["images"],         # base64 list
                        "filenames": filenames           # 确保一定是 list
                    }

                # -------- 不支持类型 --------
                else:
                    results[attachment_name] = {
                        "text": f"Unsupported file type: {file_extension}",
                        "images": [],
                        "filenames": [attachment_name]
                    }

            except Exception as e:
                # 处理单个文件的错误，不影响其他文件
                results[attachment_name] = {
                    "text": f"Error processing file: {str(e)}",
                    "images": [],
                    "filenames": [],
                    "error": str(e)
                }

        return results


# ========================
# 全局缓存与常量设置
# ========================
_token_cache = {}
_cache_max_size = 5000000

MODEL_TOKEN_LIMIT = {
    "gpt-4o": 600000,
    "gpt-4o-mini": 600000,
    "gpt-4-turbo": 600000,
    "gpt-35-turbo": 16384,
    "gpt-3.5-turbo": 16384,
    "gpt-5.2": 600000,
    "gpt-5.4mini": 600000,
}

# ========================
# 主入口函数
# ========================
def cal_tokens(username: str, attachment_names: list, deploy_model: str = "gpt-4o"):
    """
    快速计算文件的token数量 - 优化版本
    Args:
        username: 用户名
        attachment_names: 文件名列表
        deploy_model: 模型名称，默认 gpt-4o
    Returns:
        dict: {"total_tokens": int, "file_tokens": {filename: int, ...}, "limit": int}
    """
    if not isinstance(attachment_names, list):
        return {"error": "Invalid attachment_names format", "total_tokens": 0}

    try:
        try:
            # 仅当模型为 gpt-4 或 gpt-5 时忽略小版本
            if deploy_model.startswith("gpt-4") or deploy_model.startswith("gpt-5"):
                encoding = tiktoken.get_encoding("o200k_base")
            else:
                encoding = tiktoken.encoding_for_model(deploy_model)
        except Exception:
            try:
                encoding = tiktoken.get_encoding("o200k_base")
            except Exception:
                encoding = tiktoken.encoding_for_model("gpt-4")

        total_tokens = 0
        file_tokens = {}

        model_limit = MODEL_TOKEN_LIMIT.get(deploy_model, 600000)

        for attachment_name in attachment_names:
            try:
                cache_key = f"{username}:{attachment_name}:{deploy_model}"
                if cache_key in _token_cache:
                    tokens = _token_cache[cache_key]
                    file_tokens[attachment_name] = tokens
                    total_tokens += tokens
                    continue

                file_extension = attachment_name.rsplit(".", 1)[1].lower()
                blob_client = current_app.container_client.get_blob_client(f"{username}/{attachment_name}")

                tokens = _estimate_tokens_fast(blob_client, file_extension, encoding)

                # 缓存结果
                _cache_with_limit(cache_key, tokens)
                file_tokens[attachment_name] = tokens
                total_tokens += tokens

            except Exception as e:
                file_tokens[attachment_name] = 0
                print(f"⚠️ Error processing {attachment_name}: {e}")

        return {
            "total_tokens": total_tokens,
            "file_tokens": file_tokens,
            "limit": model_limit,
            "within_limit": total_tokens <= model_limit,
            "success": True,
        }

    except Exception as e:
        return {"error": str(e), "total_tokens": 0, "success": False}

# ========================
# 缓存函数
# ========================
def _cache_with_limit(key: str, value: int):
    """带上限的LRU缓存"""
    global _token_cache
    if len(_token_cache) >= _cache_max_size:
        keys_to_remove = list(_token_cache.keys())[: _cache_max_size // 4]  # 仅删1/4
        for k in keys_to_remove:
            _token_cache.pop(k, None)
    _token_cache[key] = value

# ========================
# 改进版快速估算函数
# ========================
def _estimate_tokens_fast(blob_client, file_extension: str, encoding):
    """
    针对 5MB 以内小文件优化的 Token 估算函数
    """
    try:
        blob_properties = blob_client.get_blob_properties()
        file_size = blob_properties.size  # bytes

        # === 图片类文件 ===
        if file_extension in ["jpg", "jpeg", "png"]:
            if file_size < 100 * 1024: return 100
            elif file_size < 500 * 1024: return 200
            else: return 300

        # === PDF 文件（由于限额 5MB，直接读取全量数据进行精确解析） ===
        if file_extension == "pdf":
            # 5MB 以内直接全量下载，确保解析 100% 成功
            full_data = blob_client.download_blob().readall()
            
            try:
                with fitz.open(stream=full_data, filetype="pdf") as doc:
                    n_pages = len(doc)
                    n_images = 0
                    # 快速检查前 10 页的图像密度作为采样
                    for i in range(min(10, n_pages)):
                        n_images += len(doc[i].get_images(full=True))
                    
                    img_ratio = min(n_images / max(min(10, n_pages), 1), 1.0)
            except Exception:
                # 如果 fitz 失败，回退到特征匹配
                pages_hint = full_data.count(b"/Type /Page")
                images_hint = full_data.count(b"/Subtype /Image")
                n_pages = max(1, pages_hint)
                img_ratio = min(images_hint / n_pages, 1.0) if n_pages > 0 else 0.5

            # 调优后的权重：文字页 500 tokens，图片页 300 tokens
            text_tpp = 500  
            image_tpp = 300 
            tokens_est = int(n_pages * ((1 - img_ratio) * text_tpp + img_ratio * image_tpp))
            
            # 针对 5MB 以内的 PDF，防止结构化数据导致的虚高
            return max(150, tokens_est)

        # === 文本类（TXT/JSON/CSV） ===
        if file_extension in ["txt", "json", "csv"]:
            # 5MB 以内直接读取前 1MB 采样即可
            sample_size = min(1024 * 1024, file_size)
            sample_data = blob_client.download_blob(offset=0, length=sample_size).readall()
            encoding_type = chardet.detect(sample_data)["encoding"] or "utf-8"
            sample_text = sample_data.decode(encoding_type, errors="ignore")
            
            if not sample_text: return int(file_size / 4)
            
            sample_tokens = len(encoding.encode(sample_text))
            token_density = sample_tokens / max(len(sample_text), 1)
            # 修正密度：通常 1 个字符约 0.5~0.8 token (对于 tiktoken)
            token_density = max(0.1, min(token_density, 1.2))
            
            avg_bytes_per_char = max(len(sample_data) / max(len(sample_text), 1), 1.0)
            estimated_chars = file_size / avg_bytes_per_char
            return int(estimated_chars * token_density)

        # === Excel / Word 文件 ===
        if file_extension in ["xlsx", "xls", "docx"]:
            # 5MB 以内的 Office 文件，通常包含大量 XML 结构，Token 密度较低
            # 之前 1/10 的比例依然偏高，调大分母
            divisor = 25 if file_extension == "docx" else 30
            return max(100, int(file_size / divisor))

        # === 其他未知类型 ===
        return int(file_size / 10)

    except Exception as e:
        print(f"⚠️ Error estimating tokens: {e}")
        return 100

    except Exception as e:
        print(f"⚠️ Error estimating tokens: {e}")
        return 100

# ========================
# 实用函数
# ========================
def clear_token_cache():
    global _token_cache
    _token_cache.clear()

def get_cache_stats():
    return {"cache_size": len(_token_cache), "max_size": _cache_max_size}


if __name__ == '__main__':
    file_get = FileOperation()
    # content = file_get("./", ["9Q311103_(Token：54982).pdf"])
    # print(content)
